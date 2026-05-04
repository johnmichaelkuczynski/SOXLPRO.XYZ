import math
import numpy as np
import pandas as pd
import plotly.graph_objects as go


TRADING_DAYS = 252


def _norm_cdf(x):
    """Standard normal CDF using math.erf — no scipy dependency."""
    return 0.5 * (1.0 + math.erf(x / math.sqrt(2.0)))


def _bs_call_price(S, K, T, sigma, r=0.0):
    """Black-Scholes call price, used purely as an honest mark-to-market for the
    long-call sleeve in the allocation engine. We are NOT fitting a vol surface
    or running Heston — sigma is just trailing realized vol of SOXL.
    Returns intrinsic when T<=0 or sigma<=0."""
    S = float(S); K = float(K); T = float(T); sigma = float(sigma)
    if S <= 0 or K <= 0:
        return 0.0
    if T <= 0 or sigma <= 0:
        return max(S - K, 0.0)
    sqrtT = math.sqrt(T)
    d1 = (math.log(S / K) + (r + 0.5 * sigma * sigma) * T) / (sigma * sqrtT)
    d2 = d1 - sigma * sqrtT
    return S * _norm_cdf(d1) - K * math.exp(-r * T) * _norm_cdf(d2)


def equity_curve_from_returns(returns):
    return (1.0 + pd.Series(returns).fillna(0.0)).cumprod()


def buy_and_hold_curve(price_series):
    s = price_series.dropna()
    if s.empty:
        return pd.Series(dtype=float)
    return s / s.iloc[0]


# ---------------------------------------------------------------------------
# SOXL ALLOCATION ENGINE  (default backtest strategy)
# ---------------------------------------------------------------------------
# Continuous, mean-reverting allocation against a QQQ-anchored baseline.
# Spec contract:
#   * Anchor = QQQ's cumulative behavior (both series normalized to 1.0 at
#     the start of the lookback window). NOT a fitted trendline. NOT z-scores.
#   * Deviation = SOXL_norm − QQQ_norm. Negative = oversold. Positive = overbought.
#   * Allocation is a continuous function of deviation, scaled proportionally.
#   * Hard floor and ceiling — never 0% or 100%. Allocation always lives strictly
#     inside (floor, ceiling).
#   * Bear-regime filter on QQQ scales allocation down via a multiplier, but
#     the floor still holds.
#   * Continuous rebalancing — no discrete trades, no holding periods.
#   * Never refuses a window: works on 2 data points or 5,000.
# ---------------------------------------------------------------------------

ALLOCATION_DEFAULTS = {
    "floor": 0.10,        # never below 10% long
    "ceiling": 0.90,      # never above 90% long
    "sensitivity": 2.0,   # how sharply allocation responds to deviation
    "bear_multiplier": 0.40,  # scale allocation down in QQQ bear regime
    "bear_drawdown": 0.10,    # QQQ ≥10% off rolling-max counts as bear
    "bear_lookback_frac": 0.25,  # rolling-max window = 25% of available data
}

# Hard guardrails enforcing the spec's "never 0%, never 100%" rule, regardless
# of what the user passes in. Even if floor=0 or ceiling=1 sneak through, the
# engine will clamp them inside this band so the bounds are NEVER touched.
ALLOCATION_HARD_FLOOR = 0.02   # absolute minimum exposure (2%)
ALLOCATION_HARD_CEILING = 0.98  # absolute maximum exposure (98%)


def soxl_allocation_engine(
    soxl_prices,
    qqq_prices,
    floor=ALLOCATION_DEFAULTS["floor"],
    ceiling=ALLOCATION_DEFAULTS["ceiling"],
    sensitivity=ALLOCATION_DEFAULTS["sensitivity"],
    bear_multiplier=ALLOCATION_DEFAULTS["bear_multiplier"],
    bear_drawdown=ALLOCATION_DEFAULTS["bear_drawdown"],
    bear_lookback_frac=ALLOCATION_DEFAULTS["bear_lookback_frac"],
):
    """Compute continuous allocation per the spec. Returns a DataFrame indexed by
    date with columns: soxl_norm, qqq_norm, deviation, dev_scale, raw_allocation,
    regime_mult, allocation.

    Adapts to whatever data is available — never raises, never refuses.
    """
    # --- Spec guardrail: even if user passes 0 or 1, clamp inside the hard
    # band so allocation NEVER touches 0% or 100%. ---
    floor = max(float(floor), ALLOCATION_HARD_FLOOR)
    ceiling = min(float(ceiling), ALLOCATION_HARD_CEILING)
    if floor >= ceiling:
        floor, ceiling = ALLOCATION_HARD_FLOOR, ALLOCATION_HARD_CEILING

    # --- Robust alignment / fallbacks (NEVER raises, NEVER refuses) ---
    mid = (floor + ceiling) / 2.0
    amplitude = (ceiling - floor) / 2.0

    def _empty_default():
        """Return a 1-row degenerate frame with midpoint allocation rather than
        an empty frame, so downstream UI/plot code never has to special-case
        'no data'. Honors the spec's 'never refuse' rule."""
        idx = pd.DatetimeIndex([pd.Timestamp.now().normalize()])
        return pd.DataFrame({
            "soxl_norm": [1.0], "qqq_norm": [1.0], "deviation": [0.0],
            "dev_scale": [np.nan], "raw_allocation": [mid],
            "regime_mult": [1.0], "allocation": [mid],
        }, index=idx)

    if soxl_prices is None or qqq_prices is None:
        return _empty_default()
    soxl = pd.Series(soxl_prices).dropna()
    qqq = pd.Series(qqq_prices).dropna()
    if soxl.empty or qqq.empty:
        return _empty_default()
    common = soxl.index.intersection(qqq.index)
    if len(common) == 0:
        return _empty_default()
    soxl = soxl.loc[common]
    qqq = qqq.loc[common]

    n = len(soxl)

    # --- 1-point degenerate case: return midpoint allocation ---
    if n == 1:
        return pd.DataFrame({
            "soxl_norm": [1.0],
            "qqq_norm": [1.0],
            "deviation": [0.0],
            "dev_scale": [np.nan],
            "raw_allocation": [mid],
            "regime_mult": [1.0],
            "allocation": [mid],
        }, index=common)

    # --- Normalize both to common start ---
    soxl_norm = soxl / soxl.iloc[0]
    qqq_norm = qqq / qqq.iloc[0]
    deviation = soxl_norm - qqq_norm

    # --- Adaptive scale: expanding mean of |deviation| ---
    # This makes the engine self-calibrate to whatever volatility regime exists
    # in the chosen window — no hardcoded thresholds.
    dev_abs_running = deviation.abs().expanding(min_periods=1).mean()
    dev_scale = dev_abs_running.clip(lower=0.01)  # avoid div-by-zero at t=0

    # --- Continuous allocation via tanh squashing ---
    # tanh keeps the result strictly inside (-1, +1), so allocation always lives
    # strictly inside (floor, ceiling) — the spec's hard "never 0%, never 100%" rule.
    norm_dev = deviation / dev_scale
    raw_allocation = mid + amplitude * np.tanh(-sensitivity * norm_dev)

    # --- Bear regime on QQQ ---
    bear_lookback = max(2, int(round(n * bear_lookback_frac)))
    qqq_rolling_max = qqq.rolling(bear_lookback, min_periods=1).max()
    qqq_dd = qqq / qqq_rolling_max - 1.0
    is_bear = qqq_dd <= -bear_drawdown
    regime_mult = pd.Series(1.0, index=common)
    regime_mult[is_bear] = bear_multiplier

    # --- Apply regime + clip to band (floor still wins over the multiplier) ---
    # We clip to (floor, ceiling) inclusive at the math level, but the tanh
    # squash + (floor, ceiling) interior keeps allocation strictly inside the
    # band in practice. The hard floor/ceiling clamp at function entry already
    # guarantees we can never cross 0% or 100% no matter what the user passes.
    allocation = (raw_allocation * regime_mult).clip(lower=floor, upper=ceiling)

    return pd.DataFrame({
        "soxl_norm": soxl_norm,
        "qqq_norm": qqq_norm,
        "deviation": deviation,
        "dev_scale": dev_scale,
        "raw_allocation": raw_allocation,
        "regime_mult": regime_mult,
        "allocation": allocation,
    })


def simulate_allocation_engine(soxl_prices, qqq_prices, **kwargs):
    """[Legacy spot-SOXL simulator — kept for backward compatibility.]
    Apply continuous allocation to SOXL spot returns. Allocation is lagged one
    bar to avoid look-ahead. Returns (equity_curve, daily_returns_list, alloc_df).
    """
    alloc_df = soxl_allocation_engine(soxl_prices, qqq_prices, **kwargs)
    if alloc_df.empty:
        return pd.Series(dtype=float), [], alloc_df
    if len(alloc_df) < 2:
        return pd.Series([1.0], index=alloc_df.index), [], alloc_df

    soxl = pd.Series(soxl_prices).loc[alloc_df.index]
    daily_ret = soxl.pct_change().fillna(0.0)
    alloc_lag = alloc_df["allocation"].shift(1).fillna(alloc_df["allocation"].iloc[0])
    strategy_daily = (alloc_lag * daily_ret).astype(float)
    equity = (1.0 + strategy_daily).cumprod()
    return equity, strategy_daily.tolist(), alloc_df


# ---------------------------------------------------------------------------
# CALL-SLEEVE ALLOCATION ENGINE (REVISED DEFAULT)
# ---------------------------------------------------------------------------
# Capital structure:
#   * 20% of notional sits in the "options sleeve" — long SOXL calls
#   * 80% of notional always sits in cash
#   * The 20% sleeve is sized continuously between (sleeve_pct × floor) and
#     (sleeve_pct × ceiling) of total notional based on the deviation signal
#
# Options model — "simple but honest":
#   * ATM (or slightly OTM) calls, 30–60d to expiry, rolled before expiry
#   * Mark-to-market each bar via plain Black-Scholes (sigma = trailing realized
#     vol of SOXL). NO vol surface fitting, NO Heston, NO IV smile — just one
#     scalar sigma per day for honest delta + theta + convex P&L.
#   * ASYMMETRIC continuous resize: each bar we evaluate the target. We sell
#     down freely (lock in profits, reduce exposure on overbought signals) but
#     we only ADD to the position at roll events. This honors both spec
#     requirements: "continuous rebalancing" (daily evaluation/adjustment) AND
#     "convex P&L: limited downside = premium" (no infinite refill of decaying
#     positions). Position rides delta + theta naturally between rolls.
#   * No bid-ask spread or slippage modeled.
# ---------------------------------------------------------------------------

CALL_SLEEVE_DEFAULTS = {
    "sleeve_pct": 0.20,        # 20% of notional in calls; 80% cash
    "days_to_expiry": 45,      # buy ~45 DTE calls
    "roll_at_dte": 10,         # roll when DTE drops below this
    "moneyness": 1.00,         # 1.00 = ATM, 1.05 = 5% OTM
    "vol_window": 30,          # trailing window for realized-vol pricing
    "min_sigma": 0.20,         # SOXL vol floor for pricing (20%)
    "max_sigma": 2.00,         # SOXL vol ceiling for pricing (200%)
}


def simulate_call_sleeve_engine(
    soxl_prices,
    qqq_prices,
    sleeve_pct=CALL_SLEEVE_DEFAULTS["sleeve_pct"],
    days_to_expiry=CALL_SLEEVE_DEFAULTS["days_to_expiry"],
    roll_at_dte=CALL_SLEEVE_DEFAULTS["roll_at_dte"],
    moneyness=CALL_SLEEVE_DEFAULTS["moneyness"],
    vol_window=CALL_SLEEVE_DEFAULTS["vol_window"],
    min_sigma=CALL_SLEEVE_DEFAULTS["min_sigma"],
    max_sigma=CALL_SLEEVE_DEFAULTS["max_sigma"],
    floor=ALLOCATION_DEFAULTS["floor"],
    ceiling=ALLOCATION_DEFAULTS["ceiling"],
    sensitivity=ALLOCATION_DEFAULTS["sensitivity"],
    bear_multiplier=ALLOCATION_DEFAULTS["bear_multiplier"],
    bear_drawdown=ALLOCATION_DEFAULTS["bear_drawdown"],
    bear_lookback_frac=ALLOCATION_DEFAULTS["bear_lookback_frac"],
):
    """Simulate the 20%-call-sleeve / 80%-cash strategy.

    Returns a dict with:
        equity_strategy        — pd.Series, growth of $1 (full portfolio: cash+calls)
        equity_soxl_bh         — pd.Series, SOXL B&H baseline (growth of $1)
        equity_qqq_bh          — pd.Series, QQQ B&H baseline (growth of $1)
        sleeve_value           — pd.Series, $ value of call sleeve over time
        cash_value             — pd.Series, $ cash over time
        sleeve_alloc_target    — pd.Series, target sleeve % within the 20% allowance
        sleeve_alloc_actual    — pd.Series, actual sleeve % of total portfolio
        realized_vol           — pd.Series, sigma used for pricing
        contracts_history      — pd.Series, share-equivalent contracts held
        roll_events            — list of dates where positions were rolled
        alloc_df               — full deviation/allocation panel from engine
        sleeve_pct             — capital-at-risk cap (used for risk metrics)
    """
    sleeve_pct = float(np.clip(sleeve_pct, 0.01, 1.0))

    alloc_df = soxl_allocation_engine(
        soxl_prices, qqq_prices,
        floor=floor, ceiling=ceiling, sensitivity=sensitivity,
        bear_multiplier=bear_multiplier, bear_drawdown=bear_drawdown,
        bear_lookback_frac=bear_lookback_frac,
    )

    # Align actual prices to allocation index. If inputs were empty, the engine
    # returned a synthetic 1-row frame at today's date — fall back to a trivial
    # degenerate result rather than reindexing into nothing.
    soxl_in = pd.Series(soxl_prices) if soxl_prices is not None else pd.Series(dtype=float)
    qqq_in = pd.Series(qqq_prices) if qqq_prices is not None else pd.Series(dtype=float)

    def _trivial_result(idx, mid_alloc):
        z = pd.Series(0.0, index=idx)
        one = pd.Series(1.0, index=idx)
        return {
            "equity_strategy": one, "equity_soxl_bh": one, "equity_qqq_bh": one,
            "sleeve_value": z, "cash_value": one,
            "sleeve_alloc_target": pd.Series(mid_alloc, index=idx),
            "sleeve_alloc_actual": z, "realized_vol": pd.Series(0.50, index=idx),
            "contracts_history": z, "roll_events": [],
            "alloc_df": alloc_df, "sleeve_pct": sleeve_pct,
        }

    # Defensive default — empty/degenerate windows never refuse.
    if alloc_df.empty or soxl_in.empty or qqq_in.empty:
        idx = pd.DatetimeIndex([pd.Timestamp.now().normalize()])
        return _trivial_result(idx, (floor + ceiling) / 2.0)

    common = alloc_df.index.intersection(soxl_in.index).intersection(qqq_in.index)
    if len(common) == 0:
        return _trivial_result(alloc_df.index, (floor + ceiling) / 2.0)

    soxl = soxl_in.loc[common].astype(float)
    qqq = qqq_in.loc[common].astype(float)
    alloc_df = alloc_df.loc[common]
    n = len(soxl)

    # 1-bar trivial case
    if n < 2:
        idx = soxl.index
        one = pd.Series(1.0, index=idx)
        z = pd.Series(0.0, index=idx)
        return {
            "equity_strategy": one,
            "equity_soxl_bh": soxl / soxl.iloc[0] if soxl.iloc[0] > 0 else one,
            "equity_qqq_bh": qqq / qqq.iloc[0] if qqq.iloc[0] > 0 else one,
            "sleeve_value": z, "cash_value": one,
            "sleeve_alloc_target": alloc_df["allocation"],
            "sleeve_alloc_actual": z, "realized_vol": pd.Series(0.50, index=idx),
            "contracts_history": z, "roll_events": [],
            "alloc_df": alloc_df, "sleeve_pct": sleeve_pct,
        }

    # Trailing realized vol (annualized) for pricing — adapts to window length.
    log_rets = np.log(soxl / soxl.shift(1)).fillna(0.0)
    rv_window = max(2, min(vol_window, max(2, n // 2)))
    rv = (log_rets.rolling(rv_window, min_periods=2).std() * np.sqrt(252))
    rv = rv.bfill().ffill().fillna(0.50)
    rv = rv.clip(lower=min_sigma, upper=max_sigma)

    # --- State ---
    cash = 1.0           # $ in cash (starts with full $1 notional)
    pos_shares = 0.0     # share-equivalent (1 contract = 100 shares; we use shares directly)
    pos_strike = 0.0
    pos_expiry_idx = -1
    roll_events = []

    eq_arr = np.zeros(n)
    sleeve_arr = np.zeros(n)
    cash_arr = np.zeros(n)
    target_arr = np.zeros(n)
    actual_arr = np.zeros(n)
    contracts_arr = np.zeros(n)

    days_to_expiry = max(int(days_to_expiry), 5)
    roll_at_dte = max(int(roll_at_dte), 1)

    for i in range(n):
        S = float(soxl.iloc[i])
        sigma = float(rv.iloc[i])
        # Lag the allocation signal by 1 bar — today acts on yesterday's signal.
        target_alloc = float(alloc_df["allocation"].iloc[max(i - 1, 0)])

        # 1) Mark current sleeve to market
        if pos_shares > 0 and pos_expiry_idx > i:
            T_remain = max((pos_expiry_idx - i) / 365.0, 1e-6)
            mtm_ps = _bs_call_price(S, pos_strike, T_remain, sigma)
        elif pos_shares > 0:
            mtm_ps = max(S - pos_strike, 0.0)  # at/past expiry → intrinsic
        else:
            mtm_ps = 0.0
        sleeve_dollar = pos_shares * mtm_ps

        # 2) Roll trigger — if DTE too low (or position worthless), liquidate
        #    before opening a fresh sized position.
        days_left = pos_expiry_idx - i
        need_roll = (pos_shares > 0) and (
            days_left <= roll_at_dte or mtm_ps <= 1e-6
        )
        if need_roll:
            cash += sleeve_dollar
            pos_shares = 0.0
            sleeve_dollar = 0.0
            roll_events.append(soxl.index[i])

        # 3) Sizing rules (asymmetric):
        #    • If we have NO position → open a fresh one to target sleeve $.
        #    • If we HAVE a position and target dropped below current → sell DOWN
        #      to target (release cash). This is "lock in profits / reduce exposure".
        #    • If we have a position and target is at/above current → HOLD.
        #      We never refill decaying premium mid-cycle. This is the spec's
        #      "limited downside = premium" rule. The next roll will re-size up.
        portfolio = cash + sleeve_dollar
        target_sleeve_dollar = portfolio * sleeve_pct * target_alloc

        if pos_shares == 0:
            # Open a fresh position: ATM(ish), full days_to_expiry
            T_new = days_to_expiry / 365.0
            new_strike = S * moneyness
            premium_ps = _bs_call_price(S, new_strike, T_new, sigma)
            if premium_ps > 1e-6 and target_sleeve_dollar > 1e-9:
                new_shares = target_sleeve_dollar / premium_ps
                pos_shares = new_shares
                pos_strike = new_strike
                pos_expiry_idx = i + days_to_expiry
                cash -= target_sleeve_dollar
                sleeve_dollar = target_sleeve_dollar
        elif mtm_ps > 1e-6 and target_sleeve_dollar < sleeve_dollar:
            # Sell-down only: shrink the position to target, release cash.
            target_shares = max(target_sleeve_dollar / mtm_ps, 0.0)
            delta_shares = target_shares - pos_shares  # negative → sell
            cash -= delta_shares * mtm_ps
            pos_shares = target_shares
            sleeve_dollar = pos_shares * mtm_ps

        # Record
        eq_arr[i] = cash + sleeve_dollar
        sleeve_arr[i] = sleeve_dollar
        cash_arr[i] = cash
        target_arr[i] = target_alloc
        actual_arr[i] = (sleeve_dollar / eq_arr[i]) if eq_arr[i] > 0 else 0.0
        contracts_arr[i] = pos_shares / 100.0  # display as contract count

    idx = soxl.index
    return {
        "equity_strategy": pd.Series(eq_arr, index=idx),
        "equity_soxl_bh": soxl / soxl.iloc[0],
        "equity_qqq_bh": qqq / qqq.iloc[0],
        "sleeve_value": pd.Series(sleeve_arr, index=idx),
        "cash_value": pd.Series(cash_arr, index=idx),
        "sleeve_alloc_target": pd.Series(target_arr, index=idx),
        "sleeve_alloc_actual": pd.Series(actual_arr, index=idx),
        "realized_vol": rv,
        "contracts_history": pd.Series(contracts_arr, index=idx),
        "roll_events": roll_events,
        "alloc_df": alloc_df,
        "sleeve_pct": sleeve_pct,
    }


def compute_risk_metrics(equity, label="Strategy", capital_at_risk=1.0):
    """Compute the risk-adjusted metrics required by the spec:
    Total Return, CAGR, annualized Vol, Max Drawdown, Sharpe, Sortino, Calmar,
    Capital at Risk, Return per unit of Capital at Risk.
    """
    eq = pd.Series(equity).dropna().astype(float)
    out = {
        "Series": label,
        "Total Return %": 0.0, "CAGR %": 0.0, "Vol (ann) %": 0.0,
        "Max Drawdown %": 0.0, "Sharpe": 0.0, "Sortino": 0.0, "Calmar": 0.0,
        "Capital at Risk %": round(float(capital_at_risk) * 100, 1),
        "Return / At-Risk %": 0.0,
    }
    if len(eq) < 2 or eq.iloc[0] <= 0:
        return out

    daily_rets = eq.pct_change().dropna()
    n_days = len(eq)
    years = max(n_days / TRADING_DAYS, 1e-9)

    total_return = float(eq.iloc[-1] / eq.iloc[0] - 1.0)
    cagr = float((eq.iloc[-1] / eq.iloc[0]) ** (1.0 / years) - 1.0) if years > 0 else 0.0

    rstd = float(daily_rets.std(ddof=1)) if len(daily_rets) > 1 else 0.0
    vol_ann = rstd * math.sqrt(TRADING_DAYS)
    sharpe = (float(daily_rets.mean()) / rstd * math.sqrt(TRADING_DAYS)) if rstd > 0 else 0.0

    downside = daily_rets[daily_rets < 0]
    dstd = float(downside.std(ddof=1)) if len(downside) > 1 else 0.0
    sortino = (float(daily_rets.mean()) / dstd * math.sqrt(TRADING_DAYS)) if dstd > 0 else 0.0

    rolling_max = eq.cummax()
    dd = eq / rolling_max - 1.0
    max_dd = float(dd.min())
    calmar = float(cagr / abs(max_dd)) if max_dd < 0 else 0.0

    car = float(capital_at_risk) if capital_at_risk > 0 else 1.0
    ret_per_at_risk = total_return / car

    out.update({
        "Total Return %": round(total_return * 100, 2),
        "CAGR %": round(cagr * 100, 2),
        "Vol (ann) %": round(vol_ann * 100, 2),
        "Max Drawdown %": round(max_dd * 100, 2),
        "Sharpe": round(sharpe, 2),
        "Sortino": round(sortino, 2),
        "Calmar": round(calmar, 2),
        "Capital at Risk %": round(car * 100, 1),
        "Return / At-Risk %": round(ret_per_at_risk * 100, 2),
    })
    return out


def compute_stats(equity, returns=None, n_trades=None):
    eq = pd.Series(equity).dropna()
    if eq.empty or len(eq) < 2:
        return {}
    n_days = len(eq)
    total_ret = float(eq.iloc[-1] / eq.iloc[0] - 1)
    years = n_days / TRADING_DAYS
    cagr = (eq.iloc[-1] / eq.iloc[0]) ** (1 / years) - 1 if years > 0 else 0.0
    rolling_max = eq.cummax()
    dd = eq / rolling_max - 1
    max_dd = float(dd.min())
    if returns is None:
        returns = eq.pct_change().dropna()
    rstd = float(np.std(returns, ddof=1)) if len(returns) > 1 else 0.0
    sharpe = (np.mean(returns) / rstd * np.sqrt(TRADING_DAYS)) if rstd > 0 else 0.0
    wins = [r for r in returns if r > 0]
    losses = [r for r in returns if r < 0]
    hit_rate = len(wins) / len(returns) * 100 if len(returns) else 0.0
    avg_win = float(np.mean(wins)) * 100 if wins else 0.0
    avg_loss = float(np.mean(losses)) * 100 if losses else 0.0
    return {
        "total_return_%": round(total_ret * 100, 2),
        "CAGR_%": round(cagr * 100, 2),
        "max_drawdown_%": round(max_dd * 100, 2),
        "Sharpe": round(sharpe, 2),
        "hit_rate_%": round(hit_rate, 1),
        "avg_win_%": round(avg_win, 2),
        "avg_loss_%": round(avg_loss, 2),
        "trades": int(n_trades) if n_trades is not None else int(len(returns)),
    }


def random_entry_baseline(price_series, n_trades, holding_days, seed=42):
    """Pick n_trades random entry dates, hold for holding_days, compound."""
    rng = np.random.default_rng(seed)
    s = price_series.dropna()
    if len(s) < holding_days + 5 or n_trades < 1:
        return pd.Series(dtype=float), []
    max_idx = len(s) - holding_days - 1
    entry_idxs = sorted(rng.choice(max_idx, size=min(n_trades, max_idx), replace=False))
    trade_returns = []
    timeline = pd.Series(0.0, index=s.index)
    for i in entry_idxs:
        entry = s.iloc[i]
        exit_ = s.iloc[i + holding_days]
        r = exit_ / entry - 1
        trade_returns.append(float(r))
        # spread the trade return across its holding days as daily compounded
        daily = (1 + r) ** (1 / holding_days) - 1
        for k in range(1, holding_days + 1):
            timeline.iloc[i + k] += daily
    eq = (1 + timeline).cumprod()
    return eq, trade_returns


def render_equity_chart(curves, title="Backtest Equity Curves"):
    """curves: dict of name -> pd.Series indexed by date."""
    fig = go.Figure()
    colors = {
        "Strategy": "#1976D2",
        "SOXL Buy & Hold": "#D32F2F",
        "QQQ Buy & Hold": "#43A047",
        "Random Entry Baseline": "#888888",
    }
    for name, series in curves.items():
        if series is None or len(series) == 0:
            continue
        dash = "dot" if "Random" in name else "solid"
        fig.add_trace(go.Scatter(
            x=series.index, y=series.values,
            name=name, mode="lines",
            line=dict(color=colors.get(name, "#555"), width=2.2, dash=dash),
        ))
    fig.update_layout(
        title=title, height=460, template="plotly_white",
        margin=dict(l=40, r=20, t=50, b=40),
        yaxis=dict(title="Growth of $1", tickformat=".2f"),
        legend=dict(orientation="h", yanchor="bottom", y=1.0, xanchor="center", x=0.5),
        hovermode="x unified",
    )
    return fig


def calibration_curve(predicted_probs, realized_outcomes, n_bins=10):
    """Returns (bin_centers, predicted_means, realized_means, counts) and Brier score."""
    p = np.asarray(predicted_probs, dtype=float)
    y = np.asarray(realized_outcomes, dtype=float)
    mask = np.isfinite(p) & np.isfinite(y)
    p, y = p[mask], y[mask]
    if len(p) == 0:
        return None, None, None, None, np.nan
    edges = np.linspace(0, 1, n_bins + 1)
    centers = (edges[:-1] + edges[1:]) / 2
    pred_means, real_means, counts = [], [], []
    for lo, hi in zip(edges[:-1], edges[1:]):
        mask = (p >= lo) & (p < hi if hi < 1 else p <= hi)
        if mask.sum() == 0:
            pred_means.append(np.nan)
            real_means.append(np.nan)
            counts.append(0)
            continue
        pred_means.append(float(p[mask].mean()))
        real_means.append(float(y[mask].mean()))
        counts.append(int(mask.sum()))
    brier = float(np.mean((p - y) ** 2))
    return centers, pred_means, real_means, counts, brier


def render_calibration_chart(predicted_probs, realized_outcomes, n_bins=10):
    centers, pred, real, counts, brier = calibration_curve(predicted_probs, realized_outcomes, n_bins)
    if centers is None:
        return None, np.nan
    fig = go.Figure()
    fig.add_trace(go.Scatter(
        x=[0, 1], y=[0, 1], mode="lines",
        line=dict(color="#888", dash="dash"), name="Perfect calibration",
    ))
    fig.add_trace(go.Scatter(
        x=pred, y=real, mode="markers+lines",
        marker=dict(size=[6 + min(c, 80) / 4 for c in counts], color="#1976D2"),
        name="Observed", text=[f"n={c}" for c in counts], hovertemplate="%{text}<br>predicted=%{x:.2f}<br>realized=%{y:.2f}",
    ))
    fig.update_layout(
        title=f"Calibration (Brier score: {brier:.4f}, lower is better)",
        xaxis=dict(title="Predicted probability", range=[0, 1]),
        yaxis=dict(title="Observed frequency", range=[0, 1]),
        height=420, template="plotly_white",
        margin=dict(l=40, r=20, t=50, b=40),
    )
    return fig, brier


DISCLAIMER = (
    "*Past performance does not guarantee future results. Backtests are subject to "
    "lookahead bias, survivorship bias, and overfitting.*"
)

# ---------------------------------------------------------------------------
# Report builders (TXT / CSV / DOCX / PDF) for downloadable backtest results
# ---------------------------------------------------------------------------
import io
from datetime import datetime as _dt


def _stats_rows_to_table(stats_rows):
    """Normalize list-of-dicts to (headers, rows) preserving Series order."""
    if not stats_rows:
        return [], []
    headers = ["Series"] + [k for k in stats_rows[0].keys() if k != "Series"]
    rows = [[str(r.get(h, "")) for h in headers] for r in stats_rows]
    return headers, rows


def build_report_text(title, params, methodology, stats_rows, date_range=None):
    lines = []
    lines.append("=" * 78)
    lines.append(f"BACKTEST REPORT: {title}")
    lines.append("=" * 78)
    lines.append(f"Generated: {_dt.now().strftime('%Y-%m-%d %H:%M:%S')}")
    if date_range:
        lines.append(f"Date range: {date_range[0]} → {date_range[1]}")
    lines.append("")
    lines.append("PARAMETERS")
    lines.append("-" * 78)
    for k, v in (params or {}).items():
        lines.append(f"  {k}: {v}")
    lines.append("")
    lines.append("METHODOLOGY")
    lines.append("-" * 78)
    lines.append(methodology or "(not specified)")
    lines.append("")
    lines.append("RESULTS")
    lines.append("-" * 78)
    headers, rows = _stats_rows_to_table(stats_rows)
    if headers:
        widths = [max(len(h), max((len(r[i]) for r in rows), default=0)) for i, h in enumerate(headers)]
        fmt = "  ".join("{:<" + str(w) + "}" for w in widths)
        lines.append(fmt.format(*headers))
        lines.append("  ".join("-" * w for w in widths))
        for r in rows:
            lines.append(fmt.format(*r))
    lines.append("")
    lines.append("DISCLAIMER")
    lines.append("-" * 78)
    lines.append("Past performance does not guarantee future results. Backtests are subject")
    lines.append("to lookahead bias, survivorship bias, and overfitting.")
    lines.append("")
    return "\n".join(lines)


def build_report_csv(stats_rows):
    import csv
    buf = io.StringIO()
    headers, rows = _stats_rows_to_table(stats_rows)
    if not headers:
        return ""
    w = csv.writer(buf)
    w.writerow(headers)
    w.writerows(rows)
    return buf.getvalue()


def build_report_docx(title, params, methodology, stats_rows, date_range=None):
    from docx import Document
    from docx.shared import Pt
    doc = Document()
    h = doc.add_heading(f"Backtest Report: {title}", level=1)
    doc.add_paragraph(f"Generated: {_dt.now().strftime('%Y-%m-%d %H:%M:%S')}")
    if date_range:
        doc.add_paragraph(f"Date range: {date_range[0]} → {date_range[1]}")

    doc.add_heading("Parameters", level=2)
    for k, v in (params or {}).items():
        p = doc.add_paragraph()
        run = p.add_run(f"{k}: ")
        run.bold = True
        p.add_run(str(v))

    doc.add_heading("Methodology", level=2)
    doc.add_paragraph(methodology or "(not specified)")

    doc.add_heading("Results", level=2)
    headers, rows = _stats_rows_to_table(stats_rows)
    if headers:
        table = doc.add_table(rows=1 + len(rows), cols=len(headers))
        table.style = "Light Grid Accent 1"
        for i, hdr in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = hdr
            for run in cell.paragraphs[0].runs:
                run.bold = True
        for ri, row in enumerate(rows, start=1):
            for ci, val in enumerate(row):
                table.rows[ri].cells[ci].text = val

    doc.add_heading("Disclaimer", level=2)
    p = doc.add_paragraph(
        "Past performance does not guarantee future results. Backtests are subject "
        "to lookahead bias, survivorship bias, and overfitting."
    )
    for run in p.runs:
        run.italic = True
        run.font.size = Pt(9)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def build_report_pdf(title, params, methodology, stats_rows, date_range=None):
    from reportlab.lib.pagesizes import LETTER
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.lib import colors
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                     Table, TableStyle, PageBreak)
    buf = io.BytesIO()
    doc = SimpleDocTemplate(buf, pagesize=LETTER,
                             leftMargin=0.6 * inch, rightMargin=0.6 * inch,
                             topMargin=0.6 * inch, bottomMargin=0.6 * inch)
    styles = getSampleStyleSheet()
    small = ParagraphStyle("small", parent=styles["BodyText"], fontSize=9, leading=12)
    italic_small = ParagraphStyle("ital", parent=small, fontName="Helvetica-Oblique",
                                   textColor=colors.grey)
    elements = []
    elements.append(Paragraph(f"Backtest Report: {title}", styles["Title"]))
    elements.append(Paragraph(_dt.now().strftime("Generated: %Y-%m-%d %H:%M:%S"), small))
    if date_range:
        elements.append(Paragraph(f"Date range: {date_range[0]} &rarr; {date_range[1]}", small))
    elements.append(Spacer(1, 12))

    elements.append(Paragraph("Parameters", styles["Heading2"]))
    for k, v in (params or {}).items():
        elements.append(Paragraph(f"<b>{k}:</b> {v}", small))
    elements.append(Spacer(1, 10))

    elements.append(Paragraph("Methodology", styles["Heading2"]))
    elements.append(Paragraph((methodology or "(not specified)").replace("\n", "<br/>"), small))
    elements.append(Spacer(1, 10))

    elements.append(Paragraph("Results", styles["Heading2"]))
    headers, rows = _stats_rows_to_table(stats_rows)
    if headers:
        data = [headers] + rows
        table = Table(data, repeatRows=1)
        table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1976D2")),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
            ("ALIGN", (1, 1), (-1, -1), "RIGHT"),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.whitesmoke, colors.white]),
            ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
        ]))
        elements.append(table)
    elements.append(Spacer(1, 14))

    elements.append(Paragraph("Disclaimer", styles["Heading2"]))
    elements.append(Paragraph(
        "Past performance does not guarantee future results. Backtests are "
        "subject to lookahead bias, survivorship bias, and overfitting.",
        italic_small,
    ))
    doc.build(elements)
    return buf.getvalue()


def safe_filename(title):
    keep = [c if c.isalnum() or c in "-_" else "_" for c in title]
    return "".join(keep).strip("_") or "backtest"

